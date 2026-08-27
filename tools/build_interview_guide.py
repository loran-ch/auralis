from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parents[1] / "面试资料" / "课堂助手项目面试讲解文档.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "5B6573"
PALE_BLUE = "E8EEF5"
PALE_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
DXA_WIDTH = 9360


def set_font(run, size=11, color="000000", bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=120):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_w = cell._tc.tcPr.tcW
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    run = paragraph.add_run("第 ")
    set_font(run, 9, MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)
    run = paragraph.add_run(" 页")
    set_font(run, 9, MUTED)


def add_run_text(p, text, **kwargs):
    run = p.add_run(text)
    set_font(run, **kwargs)
    return run


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    add_run_text(p, text, size={1: 16, 2: 13, 3: 12}[level], color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[level], bold=True)
    return p


def add_body(doc, text, *, bold_prefix=None):
    p = doc.add_paragraph(style="Normal")
    if bold_prefix and text.startswith(bold_prefix):
        add_run_text(p, bold_prefix, bold=True, color=INK)
        add_run_text(p, text[len(bold_prefix):])
    else:
        add_run_text(p, text)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    add_run_text(p, text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    add_run_text(p, text)
    return p


def add_callout(doc, label, text):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.10)
    p.paragraph_format.right_indent = Inches(0.10)
    shade_paragraph(p, CALLOUT)
    add_run_text(p, f"{label}  ", size=10.5, color=DARK_BLUE, bold=True)
    add_run_text(p, text, size=10.5, color=INK)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    header = table.rows[0]
    set_repeat_table_header(header)
    for cell, text in zip(header.cells, headers):
        shade(cell, PALE_GRAY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        add_run_text(p, text, size=10, color=INK, bold=True)
    for data in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, data):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_run_text(p, str(text), size=9.5)
    return table


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.167

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run_text(header, "课堂助手 | 面试讲解文档", size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(footer)
    return doc


def cover(doc):
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_text(p, "课堂助手", size=30, color=INK, bold=True)
    p.paragraph_format.space_after = Pt(8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_text(p, "项目面试讲解文档", size=18, color=BLUE, bold=True)
    p.paragraph_format.space_after = Pt(26)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_text(p, "从实时课堂记录到可追溯学习助手的全链路设计", size=12, color=MUTED)
    p.paragraph_format.space_after = Pt(46)
    rows = [
        ("项目定位", "面向课堂场景的实时转录、双语翻译、课后整理与学习问答系统"),
        ("核心技术", "FastAPI · SQLAlchemy · MySQL · WebSocket · SSE · LLM Tools · FFmpeg · Tesseract"),
        ("客户端", "浏览器静态应用 + UniApp 多端客户端"),
        ("文档用途", "面试项目介绍、技术追问与架构设计说明"),
        ("更新日期", str(date.today())),
    ]
    for label, value in rows:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(5)
        add_run_text(p, f"{label}：", size=10.5, color=INK, bold=True)
        add_run_text(p, value, size=10.5, color=MUTED)
    doc.add_page_break()


def write_content(doc):
    add_heading(doc, "一、30 秒项目介绍", 1)
    add_body(doc, "课堂助手是一个把“上课过程”沉淀为可检索学习资产的应用。它在课堂中通过 WebSocket 接收音频流，调用实时 ASR 生成字幕，并按需做上下文翻译；下课后把转录、书签、视频关键帧、OCR 结果和附件组织为课堂记录，再通过简报与工具化 Agent 帮学生回顾、定位原话、查看作业和获得学习路径。")
    add_callout(doc, "面试表达重点", "项目的核心不是“接一个大模型聊天”，而是先把课堂数据结构化、可追溯，再让模型只基于经过授权的课堂证据生成回答。")

    add_heading(doc, "二、业务问题与设计目标", 1)
    add_table(doc, ["课堂痛点", "系统设计回应"], [
        ("听课时来不及记笔记，跨语言课程理解成本高", "实时字幕与可选双语翻译；每一句都有顺序与时间偏移。"),
        ("课后回顾效率低，重点、考试提示和作业容易遗漏", "自动生成课堂简报，结合书签、关键点、考点、作业和术语。"),
        ("普通问答容易幻觉，无法证明答案来自哪节课", "课堂检索工具返回受限、带引用的证据，助手回答附可追溯引用。"),
        ("教师内容需要可共享，学生只应查看已发布课程", "课程公开状态、课主身份、已完成课次三重条件控制可见范围。"),
        ("大模型成本不可控", "先规则/检索，后按需调用模型；压缩上下文、限制工具返回、记录 Token 用量和额度。"),
    ], [3000, 6360])

    add_heading(doc, "三、总体架构与请求链路", 1)
    add_body(doc, "系统采用前后端分离但同源部署的轻量架构：FastAPI 同时提供 API 和前端静态资源，避免多端口导致的登录态、CORS 与调试复杂度。浏览器版覆盖课堂实时操作，UniApp 复用账号、数据库和 API 扩展到 H5、iOS、Android 和小程序。")
    add_table(doc, ["层次", "组件", "职责"], [
        ("客户端", "HTML/CSS/JavaScript、UniApp", "录音/录像、WebSocket 音频帧、SSE 消息渲染、课堂记录与学习交互。"),
        ("接入层", "FastAPI、CORS、中间件、静态资源", "REST API、WebSocket、SSE、请求 ID、耗时头、安全响应头与健康检查。"),
        ("领域服务", "课堂、课程、转录、翻译、简报、助手、媒体、权限", "封装业务规则，隔离路由与第三方服务。"),
        ("数据层", "SQLAlchemy、MySQL", "用户、课程、课堂、字幕、书签、简报、助手消息、媒体与审计数据。"),
        ("外部能力", "ASR、翻译 API、LLM、FFmpeg、Tesseract", "语音识别、翻译、结构化总结、媒体切帧/OCR 与片段导出。"),
    ], [1350, 2600, 5410])
    add_callout(doc, "架构取舍", "原型阶段使用 FastAPI BackgroundTasks 完成课后简报和媒体处理，降低部署复杂度；生产规模下应迁移到消息队列/独立 Worker，避免长任务占用 API 进程。")

    add_heading(doc, "四、核心数据模型：把课堂变成可检索资产", 1)
    add_body(doc, "设计上区分 Course（课程系列）与 Lecture（一次实际课堂）。Course 保存教师、语言、学期、公开状态等长期属性；Lecture 记录每一次上课的状态、时长、录音、课程关联及统计数据。这样的拆分同时支持课表推荐、教师持续授课和按次回放。")
    add_table(doc, ["实体", "关键字段", "设计意图"], [
        ("User / UserToken", "role、password_hash、access_token、refresh_token、revoked", "身份、角色与多设备会话可独立失效。"),
        ("Course", "user_id、is_public、is_active、语言、学期", "课程所有权与公开目录；归档会自动取消公开。"),
        ("Lecture", "course_id、status、duration、sentence_count、started_at", "一次课堂的生命周期和聚合指标。"),
        ("Transcription", "sentence_order、start_offset_ms、source/translated_text", "保证字幕顺序、回放定位与双语对照。"),
        ("Bookmark / Briefing", "tag、note；overview、key_points、assignments", "人工重点与自动总结并存，且支持用户修订。"),
        ("AssistantThread / Message", "lecture_ids、summary、citations", "把对话范围、上下文摘要与答案引用持久化。"),
        ("MediaAsset / ClipCandidate", "media_type、OCR 元数据、时间区间、状态", "视频、关键帧和候选片段解耦，避免自动导出浪费资源。"),
    ], [2100, 3300, 3960])
    add_body(doc, "数据库侧为常见查询建立了组合索引，例如按用户/状态/日期检索课堂、按课程/日期检索课次、按助手会话更新时间排序。字幕还使用 (lecture_id, sentence_order) 唯一约束，防止并发写入造成顺序重复。")

    add_heading(doc, "五、模块 1：认证、权限与账号体系", 1)
    add_body(doc, "认证采用 Access Token + Refresh Token 双令牌。短生命周期 Access Token 用于业务接口，Refresh Token 用于续期；服务端保存令牌会话记录并支持单设备或全设备退出，因此不是纯无状态 JWT，而是兼顾了可撤销性。")
    add_bullet(doc, "角色模型为 user、admin、super_admin：普通用户以学习者身份使用；admin 可作为教师管理并公开自己的课程；super_admin 才能进入平台级管理。")
    add_bullet(doc, "课程读取采用“课主可读 + 非课主仅可读公开且活跃课程”的规则；公开课程的访客只能查看课主已完成的课次。")
    add_bullet(doc, "请求中间件附加 X-Request-ID 和耗时头，并设置 nosniff、DENY iframe、Referrer-Policy、麦克风/摄像头 Permissions-Policy；生产环境开启 HSTS。")
    add_bullet(doc, "生产配置强制独立数据库连接、32 位以上 JWT 密钥、显式 CORS 白名单，并关闭接口文档和调试页。")

    add_heading(doc, "六、模块 2：实时录音、字幕与翻译", 1)
    add_body(doc, "实时链路使用 WebSocket，而不是频繁 HTTP 轮询。浏览器持续采集音频帧并发送到 /api/speech/{lecture_id}/stream；后端校验用户与课堂状态后，将音频桥接到百度或阿里云实时 ASR。ASR 返回 interim/final 事件，final 字幕才会落库并进入课后资产。")
    add_table(doc, ["阶段", "处理", "为什么这样设计"], [
        ("前端采集", "麦克风音频帧持续发送；可选录像/关键帧采集", "低延迟展示，同时保留媒体扩展能力。"),
        ("实时识别", "支持百度、阿里云实时 ASR；保留分片接口降级", "供应商可替换，网络或设备不兼容时仍有路径。"),
        ("断句与预览", "预览节流、短句合并、语义断句、长句强制 final", "避免“半句翻译”、DOM 无限增长和长时间无结果。"),
        ("上下文翻译", "使用最近若干句组成滑动窗口，并以稳定标记分隔当前句", "提升代词/术语连贯性，同时只返回当前句译文。"),
        ("持久化", "按 sentence_order、时间偏移保存原文与译文", "支撑回放、书签、简报、检索和引用。"),
    ], [1450, 3350, 4560])
    add_callout(doc, "异常策略", "翻译服务失败时保留原文而非丢失内容；静音/过短音频被识别为“无语音”而不是“ASR 故障”，避免前端误报整个服务不可用。")

    add_heading(doc, "七、模块 3：课堂生命周期、记录与知识卡片", 1)
    add_body(doc, "课堂具有 recording、paused、completed、failed 状态。开始课堂会校验语言和课程归属；暂停/恢复不丢失已有字幕；结束课堂后关闭实时流，并在后台启动简报与媒体处理。补录接口可以让已结束课堂重新进入 recording，同时保留已有的转录历史。")
    add_bullet(doc, "课堂记录页支持按日期、课程、状态、关键字筛选，便于快速回到某一次学习现场。")
    add_bullet(doc, "书签标签固定为重点、问题、考试、定义，既方便快速操作，也为后续候选片段和复习视图提供可解释信号。")
    add_bullet(doc, "知识卡片与回顾页把原始字幕转换为更适合复习的结构，但保留时间点/句序，避免摘要脱离原文。")

    add_heading(doc, "八、模块 4：课后简报与作业提取", 1)
    add_body(doc, "简报服务采用“双层生成”思路。第一层是本地抽取式算法：对字幕打分，生成概览、提纲、关键点、术语、考试提示、问题和作业候选；第二层在配置 LLM 时让模型对压缩后的课堂素材做结构化补全。模型输出并不直接信任，而是回填、校验并绑定原始句子引用。")
    add_table(doc, ["能力", "实现原理", "可靠性设计"], [
        ("抽取式简报", "按句子长度、关键词、书签和课堂结构评分", "无 LLM 或上游故障时仍能产出基础结果。"),
        ("LLM 增强", "输入压缩后的字幕，要求 JSON 结构返回", "解析失败回退抽取式结果；引用被回填至真实句序。"),
        ("作业/通知", "关键词与结构化简报共同识别", "标注 needs_confirmation，避免把猜测当成正式作业。"),
        ("人工编辑", "支持补充、确认、删除、编辑状态和历史载荷", "尊重教师/学生最终判断，自动结果可被修订。"),
    ], [1700, 3950, 3710])

    add_heading(doc, "九、模块 5：课程中心、教师发布与学生共享", 1)
    add_body(doc, "课程中心把“课程系列”和“每次上课”组织起来。教师（admin）创建课程、维护语言/学期/教室/配色等元数据，并可发布公开课程；学生在公开课程目录中只读浏览，看到的是该教师已完成的课堂记录及其处理结果。")
    add_number(doc, "教师创建并维护自己名下的课程；课程所有权由 user_id 固化。")
    add_number(doc, "管理员角色可把活跃课程设为 is_public；课程归档时自动取消公开，避免目录残留。")
    add_number(doc, "学生访问时通过 get_readable_course 校验公开且活跃；课程概览只返回课主 completed 的课次。")
    add_number(doc, "管理操作仍使用课主权限，学生只能浏览，不能编辑转录、简报或媒体。")
    add_callout(doc, "设计价值", "用课程公开状态实现最小可行的教师共享，而不引入复杂的班级、选课、成员表，符合当前“尽量不大改架构”的目标。未来可在 CourseMembership 上扩展邀请码、班级和精细授权。")

    add_heading(doc, "十、模块 6：课堂助手 Agent、上下文与工具", 1)
    add_body(doc, "学习助手的输入上下文不是把所有历史字幕直接拼接给模型。它首先确定会话范围（指定课次或整门课程），读取会话摘要和最近对话，随后通过工具检索必要的课堂证据。工具结果经过条数和片段长度截断，再交给模型组织最终回答。")
    add_table(doc, ["工具", "解决的问题", "成本与可信度控制"], [
        ("search_notebook", "查课堂原话、概念、考点、简报或附件标题", "最多返回有限证据；每项带 lecture、句序和时间引用。"),
        ("list_assignments", "列出作业、通知和人工上传的作业附件", "最多返回固定数量；保留待确认标记。"),
        ("breakdown_assignment", "按 L{lecture}A{index} 作业编号给学习步骤", "只给拆解与复习路径，明确禁止代写可直接提交的答案。"),
        ("get_notebook_overview", "盘点课堂数量、字幕、简报和附件", "返回聚合统计而非完整正文，适合先问“我有哪些资料”。"),
    ], [1900, 3720, 3740])
    add_body(doc, "工具选择有两条路径：一是 LLM 基于 function calling 选择工具；二是服务端关键词规则作为轻量路由和无模型降级。无论哪条路径，工具都在服务端执行，避免把数据库权限和完整课堂材料暴露给模型。")
    add_callout(doc, "提示词与上下文策略", "系统提示强调“优先依据课堂记录、引用证据、不能编造”；会话采用摘要 + 最近消息 + 当前问题，而不是无限累积历史；工具返回只保留必要片段，因此能同时降低 token、提升相关性和减少幻觉。")

    add_heading(doc, "十一、模块 7：流式输出、截图与多模态辅助", 1)
    add_body(doc, "课堂助手的流式接口使用 Server-Sent Events。前端收到 meta、tool_start、tool_result、delta、done 等事件：先展示正在检索/处理的状态，再逐段追加回答，最终以 done 事件落库并补齐引用。这比一次性等待模型完整输出更符合对话产品的反馈预期。")
    add_bullet(doc, "截图上传与 OCR 为“题目、报错、板书”提供辅助文本；它是对课堂检索的补充，不替代课堂证据。")
    add_bullet(doc, "工具调用过程以可见状态反馈给用户，降低“模型为什么这样回答”的黑箱感。")
    add_bullet(doc, "流式异常会回退到模板答案/本地分段输出，确保页面不会因单次模型请求失败而空白。")

    add_heading(doc, "十二、模块 8：视频、OCR 与课堂片段", 1)
    add_body(doc, "录像是可选能力，不影响纯录音课堂。课堂结束后，后台媒体任务读取视频资产，使用 FFmpeg 每隔固定时间抽帧并控制最大帧数/分辨率，再用 Tesseract 对关键帧进行中文+英文 OCR。识别文本、置信度和帧的时间偏移被存储，用于回顾和检索。")
    add_table(doc, ["组件", "承担的职责", "设计原因"], [
        ("FFmpeg", "音频格式转换、视频抽帧、按时间段导出片段", "成熟稳定，覆盖浏览器录制产生的常见媒体格式。"),
        ("Tesseract + 语言包", "关键帧中的板书/投影文字 OCR", "本地可运行，减少把教学画面上传给外部视觉服务的依赖。"),
        ("MediaAsset", "保存 video/frame/clip 的地址、状态和元数据", "媒体处理异步化、可观察失败状态。"),
        ("MediaClipCandidate", "基于书签/重点生成候选时间段", "先推荐、用户确认后再导出，节省 CPU 和存储。"),
    ], [1800, 3900, 3660])
    add_callout(doc, "边界说明", "当前视频理解的主体是“关键帧 OCR”，不是通用视觉大模型；它擅长识别文字型板书和课件，复杂图表/手写内容仍需要进一步引入视觉模型或人工确认。")

    add_heading(doc, "十三、可靠性、性能与安全设计", 1)
    add_table(doc, ["维度", "当前实现", "面试可主动说明的下一步"], [
        ("可用性", "翻译、LLM、ASR 均有降级/回退；健康检查区分进程存活和数据库就绪。", "引入熔断、重试退避、死信队列与第三方 SLA 监控。"),
        ("性能", "数据库连接池、索引、翻译短期缓存、字幕/工具返回上限、分页查询。", "Redis 缓存与分布式限流；Worker 横向扩容；对象存储 + CDN。"),
        ("安全", "JWT 会话撤销、角色控制、生产环境 CORS/JWT 校验、安全响应头、上传大小限制。", "令牌迁至 Secure HttpOnly Cookie；密钥管理服务；病毒扫描与审计告警。"),
        ("数据治理", "课堂证据按用户和课程范围查询；助手工具服务端执行并裁剪输出。", "增加数据保留策略、导出/删除流程、租户隔离和敏感信息脱敏。"),
        ("任务处理", "课后简报/媒体用 BackgroundTasks 异步执行。", "迁移 Celery/RQ/Arq + Redis/RabbitMQ，提供进度、重试和幂等任务 ID。"),
    ], [1300, 4000, 4060])

    add_heading(doc, "十四、前端体验与交互设计", 1)
    add_body(doc, "前端以“课堂进行中”和“课后回顾”两条主路径组织。录音页强调开始/暂停/结束的主操作、实时状态与字幕可读性；历史、知识卡片、课程、助手和个人中心使用统一导航。最近完成了一轮 UX 基础优化：共享视觉变量、统一焦点态/按钮反馈、Esc 关闭菜单、桌面侧栏固定、移动端单列布局与录音页顶部空间优化。")
    add_bullet(doc, "录音页：降低操作密度，避免移动端课程选择、标题和状态文字相互挤压。")
    add_bullet(doc, "课程中心：将课程信息与操作分层，教师创建/发布和学生浏览的权限边界更清晰。")
    add_bullet(doc, "课堂助手：把能力入口、工具反馈、流式回答和引用放在一个连续的认知流程中。")
    add_bullet(doc, "全局：键盘可见焦点、触控反馈、响应式布局和减少动态效果偏好，提升可访问性与舒适度。")

    add_heading(doc, "十五、测试与质量保障", 1)
    add_body(doc, "后端按模块编写 pytest：认证安全、翻译、音频格式、实时语音、简报、课程、助手、截图、附件、导出包、配额等均有测试文件。接口层使用 Pydantic schema 做输入校验，状态码和错误信息用于引导前端恢复。")
    add_bullet(doc, "建议面试时演示：健康检查 → 登录 → 开始课堂 → 发送音频/查看字幕 → 结束 → 打开简报/书签 → 课程发布 → 学生浏览 → Agent 检索作业。")
    add_bullet(doc, "对于流式 Agent，可展示 tool_start / tool_result / delta / done 的事件顺序，说明用户无需等待整段回复。")
    add_bullet(doc, "对于错误演示，可临时说明上游 ASR/翻译未配置时的保底行为：界面给出明确状态、保留已有课堂数据。")

    add_heading(doc, "十六、可直接用于面试的收尾总结", 1)
    add_body(doc, "我把项目设计成以“课堂证据”为中心的学习系统：前台解决实时记录和跨语言理解，后台把字幕、书签、媒体和附件结构化，课后通过简报和工具化 Agent 提升复习效率。技术上我重点处理了实时通信、异步任务、权限边界、上下文压缩、工具调用与可追溯引用；产品上则把教师发布和学生只读共享建立在尽量小的架构改动上。")
    add_callout(doc, "建议回答追问", "如果面试官问“为什么不用纯 RAG 向量库”，可回答：当前课堂数据规模和结构更适合关键词/规则检索 + 精确引用，成本低、解释性强；当课程数量和语义检索需求增长后，再将向量召回作为 search_notebook 的一个可插拔召回层。")

    add_heading(doc, "附录 A：关键接口地图", 1)
    add_table(doc, ["域", "代表接口/通道", "说明"], [
        ("认证", "/api/auth/*", "登录、刷新令牌、退出、资料与密码管理。"),
        ("课堂", "/api/lectures/*", "开始/暂停/恢复/结束、字幕、简报、附件、媒体、导出。"),
        ("实时音频", "WebSocket /api/speech/{lecture_id}/stream", "低延迟 ASR 与 final 字幕落库。"),
        ("翻译", "/api/translate", "独立文本翻译能力；实时链路内部会调用上下文翻译。"),
        ("课程", "/api/courses/*", "课程创建、更新、公开目录、课程概览和课表关联。"),
        ("学习助手", "/api/assistant/threads/*", "会话、截图、普通问答与 SSE 流式问答。"),
        ("系统", "/health/live、/health/ready", "进程存活与数据库就绪探针。"),
    ], [1600, 3300, 4460])


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = setup_document()
    cover(doc)
    write_content(doc)
    doc.core_properties.title = "课堂助手项目面试讲解文档"
    doc.core_properties.subject = "模块、实现原理与设计意图"
    doc.core_properties.author = "课堂助手项目组"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
